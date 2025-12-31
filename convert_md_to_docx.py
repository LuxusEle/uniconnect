
import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    # Simple styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Regex patterns
    h1_re = re.compile(r'^#\s+(.*)')
    h2_re = re.compile(r'^##\s+(.*)')
    h3_re = re.compile(r'^###\s+(.*)')
    img_re = re.compile(r'!\[(.*?)\]\((.*?)\)')
    table_row_re = re.compile(r'^\|(.*)\|')
    
    inside_table = False
    table = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        m_h1 = h1_re.match(line)
        if m_h1:
            doc.add_heading(m_h1.group(1), level=0)
            continue
            
        m_h2 = h2_re.match(line)
        if m_h2:
            doc.add_heading(m_h2.group(1), level=1)
            continue
            
        m_h3 = h3_re.match(line)
        if m_h3:
            doc.add_heading(m_h3.group(1), level=2)
            continue

        # Images
        m_img = img_re.search(line)
        if m_img:
            alt_text = m_img.group(1)
            img_path = m_img.group(2)
            
            # Clean path if needed (e.g. if it has query params or URL encoding)
            img_path = img_path.replace('%20', ' ')
            
            # Handle absolute paths or relative paths
            if img_path.startswith('/'):
                # Assuming stripping leading slash for local relative path if on Windows
                # But here we have relative 'documentation_assets/...' in current MD
                # or absolute 'C:/...'
                # Let's try to verify existence
                if os.path.exists(img_path.lstrip('/')):
                     img_path = img_path.lstrip('/')
                elif os.path.exists(img_path):
                     pass
                else: 
                     print(f"Warning: Image not found {img_path}")
            
            if os.path.exists(img_path):
                try:
                    # Enforce width to 5.5 inches (safe for margins)
                    doc.add_picture(img_path, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                except Exception as e:
                    doc.add_paragraph(f"[Error adding image: {alt_text}]")
                    print(f"Error adding image {img_path}: {e}")
            else:
                 doc.add_paragraph(f"[Image missing: {alt_text}]")
            continue

        # Tables (Very basic handling)
        m_table = table_row_re.match(line)
        if m_table:
            # For now, just render table rows as text to avoid complex parsing logic failure
            # Or try to implement basic table
            # If we are strictly following the MD structure:
            # We can skip the separator line |---|---|
            if '---' in line:
                continue
            
            cells = [c.strip() for c in line.strip('|').split('|')]
            
            # If not inside table, create one
            if not inside_table:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = cell_text
                inside_table = True
            else:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text
            continue
        else:
            inside_table = False

        # Regular Text (handle bullet points manually or just add paragraph)
        if line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line == '---':
            doc.add_page_break()
        else:
            # Handle standard paragraphs
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    convert_md_to_docx("university_presentation.md", "university_presentation_v5.docx")
